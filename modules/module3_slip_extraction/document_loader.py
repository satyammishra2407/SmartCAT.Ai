"""Load slip documents: PDF, Word, images."""
from __future__ import annotations

import io
from pathlib import Path

from smartcat_logging import get_logger

logger = get_logger("module3.loader")

IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp", ".webp"}
DOCX_SUFFIX = {".docx"}
PDF_SUFFIX = {".pdf"}

# Scanned slips often have only a title line from pdfplumber — always OCR below this.
SPARSE_TEXT_THRESHOLD = 200
SPARSE_CHARS_PER_PAGE = 100


def load_document(path: Path, openai_key: str | None = None) -> tuple[str, str]:
    """
    Return (full_text, extraction_method).
    extraction_method: text_pdf | ocr_pdf | vision_ocr | ocr_image | docx
    """
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix in PDF_SUFFIX:
        return _load_pdf(path, openai_key=openai_key)
    if suffix in DOCX_SUFFIX:
        return _load_docx(path), "docx"
    if suffix in IMAGE_SUFFIXES:
        return _load_image(path, openai_key=openai_key)
    raise ValueError(f"Unsupported slip format: {suffix} ({path.name})")


def text_is_sparse(text: str, page_count: int = 1) -> bool:
    clean = text.replace("\t", "").strip()
    if len(clean) < SPARSE_TEXT_THRESHOLD:
        return True
    if page_count > 1 and len(clean) / page_count < SPARSE_CHARS_PER_PAGE:
        return True
    return False


def _pdf_page_count(path: Path) -> int:
    try:
        import fitz

        return len(fitz.open(str(path)))
    except Exception:
        return 1


def _load_pdf(path: Path, openai_key: str | None = None) -> tuple[str, str]:
    import pdfplumber

    text_parts: list[str] = []
    table_parts: list[str] = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text() or ""
            if t.strip():
                text_parts.append(t)
            for table in page.extract_tables() or []:
                for row in table:
                    if not row:
                        continue
                    cells = [str(c or "").strip() for c in row if c]
                    if len(cells) >= 2:
                        table_parts.append(f"{cells[0]}\t{cells[-1]}")
                    elif len(cells) == 1:
                        table_parts.append(cells[0])

    full_text = "\n".join(text_parts).strip()
    if table_parts:
        full_text = full_text + "\n\n" + "\n".join(table_parts)

    page_count = _pdf_page_count(path)
    if not text_is_sparse(full_text, page_count):
        return full_text, "text_pdf"

    ocr_text = _ocr_pdf_pages(path)
    if ocr_text.strip() and not text_is_sparse(ocr_text, page_count):
        return ocr_text, "ocr_pdf"

    if openai_key and text_is_sparse(ocr_text or full_text, page_count):
        from modules.module3_slip_extraction.vision_ocr import ocr_pdf_with_vision

        vision_text = ocr_pdf_with_vision(path, openai_key)
        if vision_text.strip() and len(vision_text.strip()) > len((ocr_text or full_text).strip()):
            return vision_text, "vision_ocr"

    if ocr_text.strip():
        return ocr_text, "ocr_pdf"
    return full_text, "text_pdf"


def _ocr_single_page(page_index: int, path: Path, dpi: int) -> tuple[int, str]:
    try:
        import fitz
        from PIL import Image

        from modules.module3_slip_extraction.ocr_engine import image_to_text

        doc = fitz.open(str(path))
        page = doc[page_index]
        pix = page.get_pixmap(dpi=dpi)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return page_index, image_to_text(img)
    except Exception as e:
        logger.debug("Page OCR failed page=%s: %s", page_index, e)
        return page_index, ""


def _ocr_pdf_pages(path: Path, dpi: int = 150) -> str:
    try:
        import fitz
    except ImportError:
        logger.warning("PyMuPDF not installed — cannot OCR scanned PDF %s", path.name)
        return ""

    from modules.module3_slip_extraction.ocr_engine import tesseract_available

    if not tesseract_available():
        logger.info("Tesseract not available — skipping local OCR for %s", path.name)
        return ""

    page_count = _pdf_page_count(path)
    if page_count <= 0:
        return ""

    from concurrent.futures import ThreadPoolExecutor, as_completed

    indexed: dict[int, str] = {}
    workers = min(4, page_count)
    with ThreadPoolExecutor(max_workers=workers) as pool:
        futures = [pool.submit(_ocr_single_page, i, path, dpi) for i in range(page_count)]
        for fut in as_completed(futures):
            idx, text = fut.result()
            if text.strip():
                indexed[idx] = text

    if indexed:
        return "\n\n".join(indexed[i] for i in sorted(indexed))
    return ""


def _load_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if len(cells) >= 2:
                parts.append(f"{cells[0]}\t{cells[-1]}")
            elif cells:
                parts.append(cells[0])
    return "\n".join(parts)


def _load_image(path: Path, openai_key: str | None = None) -> str:
    from PIL import Image

    from modules.module3_slip_extraction.ocr_engine import image_to_text, tesseract_available

    img = Image.open(path)
    text = image_to_text(img)
    if text.strip() or not openai_key:
        return text

    try:
        import base64
        import io as _io

        from openai import OpenAI

        buf = _io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.standard_b64encode(buf.getvalue()).decode("ascii")
        client = OpenAI(api_key=openai_key)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Transcribe all text from this insurance slip image."},
                        {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}", "detail": "high"}},
                    ],
                }
            ],
            temperature=0,
            max_tokens=4096,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        logger.warning("Image vision OCR failed: %s", e)
        return text
